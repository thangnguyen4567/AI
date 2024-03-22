from docx import Document
from langchain.schema import Document as LangchainDocument
from config.config_vectordb import VectorDB

def split_docx_by_heading_lv1(paragraphs,doc,role):
    headingarr = [p.text for p in doc.paragraphs if p.style.name == 'Heading 1']
    docs = []
    skip = True
    # Chia tệp thành các phần dựa trên các tiêu đề cấp 2
    for i, heading in enumerate(headingarr):
        if i == 0:
            content = heading + ':\n'
            # Thêm các đoạn văn bản cho phần này
            for paragraph in paragraphs:
                if i + 1 < len(headingarr) and paragraph.text == headingarr[i+1]:
                    break
                # Bỏ qua phần đầu, khi nào đến vị trí heading mới bắt đầu lấy dữ liệu
                if paragraph.text == heading:
                    skip = False
                if skip == False:
                    content += '\n'+paragraph.text

            docs.append(LangchainDocument(page_content=content,metadata={'role':role,'title':headingarr[i]}))
    return docs

def split_docx_by_heading_lv2(paragraphs,doc,role):
    skipheader = ['Xây dựng nội dung lớp học','Xây dựng nội dung lớp học – CM','Quản lý báo cáo – RM']
    headingarr = [p.text for p in doc.paragraphs if p.style.name == 'Heading 2']
    del headingarr[0:3]
    skip = True
    docs = []
    # Chia tệp thành các phần dựa trên các tiêu đề cấp 2
    for i, heading in enumerate(headingarr):
        content = heading + ':\n'
        # Thêm các đoạn văn bản cho phần này
        for paragraph in paragraphs:
            if i + 1 < len(headingarr) and paragraph.text == headingarr[i+1]:
                break
            if heading in skipheader:
                skip = True
                content = ''
                break
            # Bỏ qua phần đầu, khi nào đến vị trí heading mới bắt đầu lấy dữ liệu
            if paragraph.text == heading:
                skip = False
  
            if skip == False:
                content += '\n'+paragraph.text

        if content != '':
            docs.append(LangchainDocument(page_content=content,metadata={'role':role,'title':headingarr[i]}))
    return docs

def split_docx_by_heading_lv3(paragraphs,doc,role):
    skip = True
    headingarr = [p.text for p in doc.paragraphs if p.style.name == 'Heading 3']
    if role == 'teacher':
        # tài liệu giảng viên
        del headingarr[0:10]
        del headingarr[16:26]
    elif role == 'manager':
        # tài liệu quản lý
        del headingarr[0:16]
        del headingarr[18:56]
    else:
        return []
    docs = []
    # Chia tệp thành các phần dựa trên các tiêu đề cấp 2
    for i, heading in enumerate(headingarr):
        content = heading + ':\n'
        # Thêm các đoạn văn bản cho phần này
        for paragraph in paragraphs:
            if i + 1 < len(headingarr) and paragraph.text == headingarr[i+1]:
                break
            if paragraph.style.name == 'Heading 2' and skip == False:
                skip = True
                break
            # Bỏ qua phần đầu, khi nào đến vị trí heading mới bắt đầu lấy dữ liệu
            if paragraph.text == heading:
                skip = False

            if skip == False:
                content += '\n'+paragraph.text

        docs.append(LangchainDocument(page_content=content,metadata={'role':role,'title':headingarr[i]}))
    return docs

# role = 'manager'
# input_file = ".\document\TAILIEU_HDSD_QLDT_LMS4.0.docx"
# role = 'teacher'
# input_file = ".\document\TAILIEU_HDSD_GV_LMS4.0.docx"
role = 'admin'
input_file = ".\document\TAILIEU_HDSD_QTV_LMS4.0.docx"
# role = 'student'
# input_file = ".\document\TAILIEU_HDSD_HV_LMS4.0.docx"
doc = Document(input_file)
paragraphs = iter(doc.paragraphs)
docslv1 = split_docx_by_heading_lv1(paragraphs,doc,role)
paragraphs = iter(doc.paragraphs)
docslv2 = split_docx_by_heading_lv2(paragraphs,doc,role)
paragraphs = iter(doc.paragraphs)
docslv3 = split_docx_by_heading_lv3(paragraphs,doc,role)
docs = docslv1 + docslv2 + docslv3
VectorDB().add_vectordb(docs,'system')